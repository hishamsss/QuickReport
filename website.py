import streamlit as st
import pandas as pd
from docx import Document
import re
from io import BytesIO
from copy import deepcopy
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path

# === Word Processing Functions ===
def classify(percentile):
    try:
        if isinstance(percentile, str) and ">" in percentile:
            percentile = float(percentile.replace(">", ""))
        elif percentile == "-" or pd.isna(percentile):
            return "-"
        else:
            percentile = float(percentile)
    except:
        return "-"

    if percentile <= 1:
        return "Extremely Low"
    elif 2 <= percentile <= 8:
        return "Unusually Low"
    elif 9 <= percentile <= 24:
        return "Low Average"
    elif 25 <= percentile <= 74:
        return "Average"
    elif 75 <= percentile <= 90:
        return "High Average"
    elif 91 <= percentile <= 97:
        return "Unusually High"
    elif percentile >= 98:
        return "Extremely High"
    else:
        return "-"

def format_percentile_with_suffix(percentile):
    try:
        if isinstance(percentile, str) and ">" in percentile:
            percentile = float(percentile.replace(">", ""))
        elif percentile == "-" or pd.isna(percentile):
            return "-"
        else:
            percentile = float(percentile)
    except:
        return "-"

    if percentile.is_integer():
        integer_part = int(percentile)
    else:
        decimal_first_digit = int(str(percentile).split(".")[1][0])
        integer_part = decimal_first_digit

    if 10 <= integer_part % 100 <= 20:
        suffix = 'th'
    else:
        last_digit = integer_part % 10
        if last_digit == 1:
            suffix = 'st'
        elif last_digit == 2:
            suffix = 'nd'
        elif last_digit == 3:
            suffix = 'rd'
        else:
            suffix = 'th'

    if percentile.is_integer():
        return f"{int(percentile)}{suffix}"
    else:
        return f"{percentile}{suffix}"

def replace_placeholders(doc, lookup):
    def replace_in_runs(runs, lookup):
        n = len(runs)
        i = 0
        while i < n:
            found = False
            combined_text = ''
            for j in range(i, n):
                combined_text += runs[j].text
                for key, value in lookup.items():
                    placeholder = f"{{{{{key}}}}}"
                    if placeholder in combined_text:
                        new_text = combined_text.replace(placeholder, value)
                        idx = 0
                        for k in range(i, j + 1):
                            run_length = len(runs[k].text)
                            runs[k].text = new_text[idx:idx+run_length]
                            idx += run_length
                        found = True
                        break
                if found:
                    break
            i = j + 1 if found else i + 1

    for para in doc.paragraphs:
        replace_in_runs(para.runs, lookup)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_in_runs(para.runs, lookup)

def superscript_suffixes(doc):
    pattern = re.compile(r'(\d+(?:\.\d+)?)(st|nd|rd|th)')

    def copy_font_settings(source_run, target_run):
        target_run.font.bold = source_run.font.bold
        target_run.font.italic = source_run.font.italic
        target_run.font.underline = source_run.font.underline
        target_run.font.size = source_run.font.size
        target_run.font.name = source_run.font.name
        if source_run.font.color.rgb:
            target_run.font.color.rgb = source_run.font.color.rgb

    def process_runs(paragraph):
        new_runs = []
        for run in paragraph.runs:
            text = run.text
            last_end = 0
            matches = list(pattern.finditer(text))

            if not matches:
                new_runs.append((text, False, run))
            else:
                for match in matches:
                    start, end = match.span()
                    if start > last_end:
                        new_runs.append((text[last_end:start], False, run))
                    new_runs.append((match.group(1), False, run))
                    new_runs.append((match.group(2), True, run))
                    last_end = end
                if last_end < len(text):
                    new_runs.append((text[last_end:], False, run))

        for run in paragraph.runs:
            run.text = ''

        for text, is_super, original_run in new_runs:
            if text == '':
                continue
            new_run = paragraph.add_run(text)
            copy_font_settings(original_run, new_run)
            if is_super:
                new_run.font.superscript = True

    for para in doc.paragraphs:
        process_runs(para)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    process_runs(para)

def delete_rows_with_dash(doc):
    for table in doc.tables:
        rows_to_delete = []
        for row_idx, row in enumerate(table.rows):
            for cell in row.cells:
                if cell.text.strip() == "#":
                    rows_to_delete.append(row_idx)
                    break
        for row_idx in sorted(rows_to_delete, reverse=True):
            tbl = table._tbl
            tr = table.rows[row_idx]._tr
            tbl.remove(tr)

def delete_paragraphs_containing_dash(doc):
    def delete_paragraph(paragraph):
        p = paragraph._element
        parent = p.getparent()
        if parent is not None:
            parent.remove(p)

    paras_to_delete = [para for para in doc.paragraphs if "-" in para.text]
    for para in paras_to_delete:
        delete_paragraph(para)

# === Streamlit App ===
st.title("📄 Word Document Filler")

uploaded_doc = st.file_uploader("Upload Your WIAT-4 Report (.docx)", type="docx")
TEMPLATE_PATH = Path("template.docx")
template_doc = Document(TEMPLATE_PATH)

uploaded_doc = st.file_uploader("Upload Input Word Document (.docx)", type="docx")

if uploaded_doc:
    if st.button("Generate Filled Template"):
        input_doc = Document(uploaded_doc)
        working_template = Document(TEMPLATE_PATH)

        # Process tables
        target_table_indices = [2, 4, 9]
        ae_combined = pd.DataFrame()

        for i in target_table_indices:
            if i < len(input_doc.tables):
                table = input_doc.tables[i]
                data = []
                for row in table.rows:
                    data.append([cell.text.strip() for cell in row.cells])
                df = pd.DataFrame(data)

                if df.shape[0] > 1:
                    df.columns = df.iloc[0]
                    df = df.drop(index=0).reset_index(drop=True)

                if df.shape[1] >= 5:
                    ae_df = df.iloc[:, [0, 4]].copy()
                    ae_df.columns = ['Name', 'Percentile']
                    ae_combined = pd.concat([ae_combined, ae_df], ignore_index=True)

        if not ae_combined.empty:
            ae_combined.drop_duplicates(subset='Name', inplace=True)
            ae_combined["Classification"] = ae_combined["Percentile"].apply(classify)
            ae_combined["Percentile*"] = ae_combined["Percentile"].apply(format_percentile_with_suffix)
            ae_combined = ae_combined.applymap(lambda x: "#" if x == "-" else x)

        # Build lookup
        lookup = {}
        for idx, row in ae_combined.iterrows():
            name = row['Name'].strip()
            lookup[f"{name} Classification"] = row['Classification']
            lookup[f"{name} Percentile"] = str(row['Percentile']).strip()
            lookup[f"{name} Percentile*"] = str(row['Percentile*']).strip()

        replace_placeholders(template_doc, lookup)

        superscript_suffixes(template_doc)

        delete_rows_with_dash(template_doc)
        #delete_paragraphs_containing_dash(template_doc)

        # Save to BytesIO
        output = BytesIO()
        template_doc.save(output)
        output.seek(0)

        st.success("✅ Document generated successfully!")

        st.download_button(
            label="📥 Download Filled Template",
            data=output,
            file_name="filled_template.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
