import streamlit as st
import pandas as pd
import re
from docx import Document
from io import BytesIO
from docx.enum.text import WD_COLOR_INDEX

# === Helper Functions ===

def classify(percentile):
    try:
        if isinstance(percentile, str) and ">" in percentile:
            percentile = float(percentile.replace(">", ""))
        elif percentile == "-" or pd.isna(percentile):
            return "-"
        else:
            percentile = float(percentile)
    except:
        return "-"

    if percentile <= 1:
        return "Extremely Low"
    elif 2 <= percentile <= 8:
        return "Unusually Low"
    elif 9 <= percentile <= 24:
        return "Low Average"
    elif 25 <= percentile <= 74:
        return "Average"
    elif 75 <= percentile <= 90:
        return "High Average"
    elif 91 <= percentile <= 97:
        return "Unusually High"
    elif percentile >= 98:
        return "Extremely High"
    else:
        return "-"

def format_percentile_with_suffix(percentile):
    try:
        if isinstance(percentile, str) and ">" in percentile:
            percentile = float(percentile.replace(">", ""))
        elif percentile == "-" or pd.isna(percentile):
            return "-"
        else:
            percentile = float(percentile)
    except:
        return "-"

    if percentile.is_integer():
        integer_part = int(percentile)
    else:
        decimal_first_digit = int(str(percentile).split(".")[1][0])
        integer_part = decimal_first_digit

    if 10 <= integer_part % 100 <= 20:
        suffix = 'th'
    else:
        last_digit = integer_part % 10
        if last_digit == 1:
            suffix = 'st'
        elif last_digit == 2:
            suffix = 'nd'
        elif last_digit == 3:
            suffix = 'rd'
        else:
            suffix = 'th'

    if percentile.is_integer():
        return f"{int(percentile)}{suffix}"
    else:
        return f"{percentile}{suffix}"

def replace_placeholders(doc, lookup):
    def replace_in_runs(runs, lookup):
        n = len(runs)
        i = 0
        while i < n:
            found = False
            combined_text = ''
            for j in range(i, n):
                combined_text += runs[j].text
                for key, value in lookup.items():
                    placeholder = f"{{{{{key}}}}}"
                    if placeholder in combined_text:
                        new_text = combined_text.replace(placeholder, value)
                        idx = 0
                        for k in range(i, j + 1):
                            run_length = len(runs[k].text)
                            runs[k].text = new_text[idx:idx+run_length]
                            idx += run_length
                        found = True
                        break
                if found:
                    break
            i = j + 1 if found else i + 1

    for para in doc.paragraphs:
        replace_in_runs(para.runs, lookup)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_in_runs(para.runs, lookup)

def superscript_suffixes(doc):
    pattern = re.compile(r'(\d+(?:\.\d+)?)(st|nd|rd|th)')

    def copy_font_settings(source_run, target_run):
        target_run.font.bold = source_run.font.bold
        target_run.font.italic = source_run.font.italic
        target_run.font.underline = source_run.font.underline
        target_run.font.size = source_run.font.size
        target_run.font.name = source_run.font.name
        if source_run.font.color.rgb:
            target_run.font.color.rgb = source_run.font.color.rgb

    def process_runs(paragraph):
        new_runs = []
        for run in paragraph.runs:
            text = run.text
            last_end = 0
            matches = list(pattern.finditer(text))

            if not matches:
                new_runs.append((text, False, run))
            else:
                for match in matches:
                    start, end = match.span()
                    if start > last_end:
                        new_runs.append((text[last_end:start], False, run))
                    new_runs.append((match.group(1), False, run))
                    new_runs.append((match.group(2), True, run))
                    last_end = end
                if last_end < len(text):
                    new_runs.append((text[last_end:], False, run))

        for run in paragraph.runs:
            run.text = ''

        for text, is_super, original_run in new_runs:
            if text == '':
                continue
            new_run = paragraph.add_run(text)
            copy_font_settings(original_run, new_run)
            if is_super:
                new_run.font.superscript = True

    for para in doc.paragraphs:
        process_runs(para)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    process_runs(para)

def delete_rows_with_dash(doc):
    for table in doc.tables:
        rows_to_delete = []
        for row_idx, row in enumerate(table.rows):
            for cell in row.cells:
                if cell.text.strip() == "#":
                    rows_to_delete.append(row_idx)
                    break
        for row_idx in sorted(rows_to_delete, reverse=True):
            tbl = table._tbl
            tr = table.rows[row_idx]._tr
            tbl.remove(tr)

def delete_rows_with_unfilled_placeholders(doc):
    pattern = re.compile(r"\{\{.*?\}\}")
    for table in doc.tables:
        rows_to_delete = []
        for row_idx, row in enumerate(table.rows):
            for cell in row.cells:
                if pattern.search(cell.text):
                    rows_to_delete.append(row_idx)
                    break
        for row_idx in sorted(rows_to_delete, reverse=True):
            tbl = table._tbl
            tr = table.rows[row_idx]._tr
            tbl.remove(tr)

def highlight_unfilled_placeholders(doc):
    placeholder_pattern = re.compile(r"\{\{.*?\}\}")
    missing_symbol_pattern = re.compile(r"#")

    def highlight_placeholder_in_runs(runs):
        combined_text = ''
        run_indices = []

        for idx, run in enumerate(runs):
            combined_text += run.text
            run_indices.append(idx)

        matches = list(placeholder_pattern.finditer(combined_text)) + list(missing_symbol_pattern.finditer(combined_text))
        if not matches:
            return

        current_pos = 0
        for idx in run_indices:
            run = runs[idx]
            text_len = len(run.text)
            run_end_pos = current_pos + text_len

            for match in matches:
                match_start, match_end = match.span()
                if match_start < run_end_pos and match_end > current_pos:
                    run.font.highlight_color = WD_COLOR_INDEX.YELLOW

            current_pos = run_end_pos

    for para in doc.paragraphs:
        highlight_placeholder_in_runs(para.runs)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    highlight_placeholder_in_runs(para.runs)

# === Streamlit App ===

st.title("\U0001F4C4 Report Writer")

tab1, tab2, tab3, tab4, tab5 = st.tabs(["WIAT", "WISC", "ChAMP", "Beery", "Finalize"])

with tab1:
    uploaded_doc = st.file_uploader("\U0001F4C4 Upload WIAT-4 Report (.docx)", type="docx", key="wiat_upload")

with tab2:
    uploaded_wisc = st.file_uploader("🧠 Upload WISC Report (.docx)", type="docx", key="wisc_upload")

with tab3:
    st.subheader("🧠 Enter ChAMP Scores")

    champ_fields = [
        "Lists", "Objects", "Instructions", "Places", "Lists Delayed",
        "Lists Recognition", "Objects Delayed", "Instructions Delayed",
        "Instructions Recognition", "Places Delayed", "Verbal Memory Index",
        "Visual Memory Index", "Immediate Memory Index", "Delayed Memory Index",
        "Total Memory Index", "Screening Index"
    ]

    champ_data = []
    for field in champ_fields:
        col1, col2 = st.columns([1.5, 1.5])
        with col1:
            st.markdown(f"**{field}**")
        with col2:
            value = st.text_input("", key=f"champ_{field}")
        champ_data.append({"Name": field, "Percentile": value})

    champ_df = pd.DataFrame(champ_data)
    if not champ_df["Percentile"].eq("").all():
        champ_df["Classification"] = champ_df["Percentile"].apply(classify)
        champ_df["Percentile*"] = champ_df["Percentile"].apply(format_percentile_with_suffix)
        champ_df = champ_df.replace("-", "#")

with tab4:
    st.subheader("✍️ Enter Beery Scores")

    col1, col2 = st.columns(2)
    with col1:
        vmi = st.text_input("Visual-Motor Integration (VMI) Percentile", key="vmi_input")
    with col2:
        vmi_raw = st.text_input("VMI Raw Score", key="vmi_raw_input")

    col1, col2 = st.columns(2)
    with col1:
        vp = st.text_input("Visual Perception (VP) Percentile", key="vp_input")
    with col2:
        vp_raw = st.text_input("VP Raw Score", key="vp_raw_input")

    col1, col2 = st.columns(2)
    with col1:
        mc = st.text_input("Motor Coordination (MC) Percentile", key="mc_input")
    with col2:
        mc_raw = st.text_input("MC Raw Score", key="mc_raw_input")

with tab5:
    st.subheader("Report Settings")

    # 1) Always-visible fields:
    report_name_input = st.text_input(
        "Report file name (without .docx)",
        value="combined_report",
        key="report_name_input"
    )  
    gender_selection = st.radio(
        "Select WIAT Report Gender Template:",
        ("Male", "Female"),
        key="gender"
    )

    # 2) If files aren’t uploaded yet, prompt the user:
    if not uploaded_doc or not uploaded_wisc:
        st.info("Please upload both your WIAT and WISC reports in the WIAT & WISC tabs above.")
    else:
        # 3) Once both are present, show the generate button
        if st.button("Generate Combined Report"):
            # … your existing document-generation logic here …
            input_doc    = Document(uploaded_doc)
            template_path = (
                "template_male.docx"
                if gender_selection == "Male"
                else "template_female.docx"
            )
            template_doc = Document(template_path)
        
            # 2) (Now your existing AE‐table loops, placeholder‐replacing, superscripting, etc.)
        
            # 3) Save into bytes
            output = BytesIO()
            template_doc.save(output)
            st.session_state["generated_report"] = output.getvalue()            
            st.success("✅ Combined document generated successfully!")

            # === Process WIAT Tables ===
            ae_combined = pd.DataFrame()

            for i, table in enumerate(input_doc.tables):
                data = [[cell.text.strip() for cell in row.cells] for row in table.rows]
                df = pd.DataFrame(data)
                if df.shape[0] > 1:
                    df.columns = df.iloc[0]
                    df = df.drop(index=0).reset_index(drop=True)
                if df.shape[1] >= 5:
                    ae_df = df.iloc[:, [0, 4]].copy()
                    ae_df.columns = ['Name', 'Percentile']
                    ae_df['Name'] = ae_df['Name'].str.replace(r'[^A-Za-z\s]', '', regex=True).str.strip()
                    ae_combined = pd.concat([ae_combined, ae_df], ignore_index=True)
                    
            if not ae_combined.empty:
                ae_combined.drop_duplicates(subset='Name', inplace=True)
                ae_combined["Classification"] = ae_combined["Percentile"].apply(classify)
                ae_combined["Percentile*"] = ae_combined["Percentile"].apply(format_percentile_with_suffix)
                ae_combined = ae_combined.replace("-", "#")

            lookup = {}
            for _, row in ae_combined.iterrows():
                name = row['Name'].strip()
                lookup[f"{name} Classification"] = row['Classification']
                lookup[f"{name} Percentile"] = str(row['Percentile']).strip()
                lookup[f"{name} Percentile*"] = str(row['Percentile*']).strip()

            # === Beery
            if vmi:
                lookup["VMI Percentile"] = vmi
                lookup["VMI Percentile*"] = format_percentile_with_suffix(vmi)
                lookup["VMI Classification"] = classify(vmi)
            if vmi_raw:
                lookup["VMI Raw Score"] = vmi_raw
            if vp:
                lookup["VP Percentile"] = vp
                lookup["VP Percentile*"] = format_percentile_with_suffix(vp)
                lookup["VP Classification"] = classify(vp)
            if vp_raw:
                lookup["VP Raw Score"] = vp_raw
            if mc:
                lookup["MC Percentile"] = mc
                lookup["MC Percentile*"] = format_percentile_with_suffix(mc)
                lookup["MC Classification"] = classify(mc)
            if mc_raw:
                lookup["MC Raw Score"] = mc_raw

            # === ChAMP
            if not champ_df.empty:
                for _, row in champ_df.iterrows():
                    name = row['Name'].strip()
                    lookup[f"{name} Percentile"] = row['Percentile']
                    lookup[f"{name} Percentile*"] = row['Percentile*']
                    lookup[f"{name} Classification"] = row['Classification']

            # === WISC
            input_wisc_doc = Document(uploaded_wisc)
            wisc_combined = pd.DataFrame()

            for i, table in enumerate(input_wisc_doc.tables):
                data = [[cell.text.strip() for cell in row.cells] for row in table.rows]
                df = pd.DataFrame(data)
                if df.shape[0] > 1:
                    df.columns = df.iloc[0]
                    df = df.drop(index=0).reset_index(drop=True)
                    if df.shape[1] >= 5:
                        if i == 5 or i == 15:
                            ae_df = df.iloc[:, [1, 4]].copy()
                        elif df.shape[1] >= 6:
                            ae_df = df.iloc[:, [1, 5]].copy()
                        else:
                            continue
                        ae_df.columns = ['Name', 'Percentile']
                        ae_df['Name'] = ae_df['Name'].str.replace(r'[^A-Za-z\s]', '', regex=True).str.strip()
                        wisc_combined = pd.concat([wisc_combined, ae_df], ignore_index=True)

            if not wisc_combined.empty:
                wisc_combined.drop_duplicates(subset='Name', inplace=True)
                wisc_combined["Classification"] = wisc_combined["Percentile"].apply(classify)
                wisc_combined["Percentile*"] = wisc_combined["Percentile"].apply(format_percentile_with_suffix)
                wisc_combined = wisc_combined.replace("-", "#")

                for _, row in wisc_combined.iterrows():
                    name = row['Name'].strip()
                    lookup[f"{name} Classification"] = row['Classification']
                    lookup[f"{name} Percentile"] = str(row['Percentile']).strip()
                    lookup[f"{name} Percentile*"] = str(row['Percentile*']).strip()
            

            # === Fill and output unified report
            replace_placeholders(template_doc, lookup)
            superscript_suffixes(template_doc)
            delete_rows_with_dash(template_doc)
            delete_rows_with_unfilled_placeholders(template_doc)
            highlight_unfilled_placeholders(template_doc)

            output = BytesIO()
            template_doc.save(output)
            st.session_state["generated_report"] = output.getvalue()
            

        if st.session_state.get("generated_report"):
            output_data = BytesIO(st.session_state["generated_report"])
            final_name = report_name_input.strip() or "combined_report"
            if not final_name.lower().endswith(".docx"):
                final_name += ".docx"

            st.download_button(
                label="📥 Download Combined Report",
                data=output_data,
                file_name=final_name,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
